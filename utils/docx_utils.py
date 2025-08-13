import io
from typing import Dict, Any, List
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DOCXProcessor:
    """Process DOCX documents and extract text."""
    
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            with io.BytesIO(file_bytes) as buf:
                doc = Document(buf)
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                # Extract text from tables
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            tables.append(" | ".join(row_text))
                
                # Combine all text
                all_text = []
                if paragraphs:
                    all_text.extend(paragraphs)
                if tables:
                    all_text.extend(tables)
                
                return "\n\n".join(all_text).strip()
                
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return ""
    
    @staticmethod
    def get_document_info(file_bytes: bytes) -> Dict[str, Any]:
        """Get DOCX metadata."""
        try:
            with io.BytesIO(file_bytes) as buf:
                doc = Document(buf)
                
                # Count paragraphs and tables
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                table_count = len(doc.tables)
                
                return {
                    "paragraphs": paragraph_count,
                    "tables": table_count,
                    "size_kb": len(file_bytes) / 1024,
                    "core_properties": {
                        "title": doc.core_properties.title or "Unknown",
                        "author": doc.core_properties.author or "Unknown",
                        "created": str(doc.core_properties.created) if doc.core_properties.created else "Unknown",
                        "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "Unknown"
                    }
                }
                
        except Exception as e:
            logger.error(f"Error getting DOCX info: {e}")
            return {
                "paragraphs": 0,
                "tables": 0,
                "size_kb": 0,
                "core_properties": {}
            }
    
    @staticmethod
    async def extract_with_metadata(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text and metadata together."""
        info = DOCXProcessor.get_document_info(file_bytes)
        text = DOCXProcessor.extract_text(file_bytes)
        
        return {
            "text": text,
            "metadata": info,
            "document_type": DOCXProcessor.classify_document_type(text)
        }
    
    @staticmethod
    def classify_document_type(text: str) -> str:
        """Classify document type based on content."""
        text_lower = text.lower()
        
        # English classifications
        if any(word in text_lower for word in ['contract', 'agreement', 'terms']):
            return "Contract/Agreement"
        elif any(word in text_lower for word in ['invoice', 'bill', 'payment']):
            return "Invoice/Bill"
        elif any(word in text_lower for word in ['resume', 'cv', 'curriculum vitae']):
            return "Resume/CV"
        elif any(word in text_lower for word in ['report', 'analysis', 'summary']):
            return "Report/Analysis"
        elif any(word in text_lower for word in ['policy', 'procedure', 'guideline']):
            return "Policy/Procedure"
        
        # Tamil classifications
        elif any(word in text_lower for word in ['ஒப்பந்தம்', 'ஒப்பந்த', 'உடன்படிக்கை']):
            return "ஒப்பந்தம்/உடன்படிக்கை"
        elif any(word in text_lower for word in ['இன்வாய்ஸ்', 'பில்', 'கட்டணம்']):
            return "இன்வாய்ஸ்/பில்"
        elif any(word in text_lower for word in ['ரெசுமே', 'சிவி', 'வேலை விண்ணப்பம்']):
            return "ரெசுமே/சிவி"
        elif any(word in text_lower for word in ['அறிக்கை', 'பகுப்பாய்வு', 'சுருக்கம்']):
            return "அறிக்கை/பகுப்பாய்வு"
        
        return "General Document"
